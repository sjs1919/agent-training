"""将 8周执行计划-v5.md 和 每日打卡表-v5.md 合并生成 Word 文档。"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

BASE = Path(__file__).resolve().parent

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_markdown_to_doc(doc, text):
    """Parse markdown text and add to doc with basic formatting."""
    lines = text.split('\n')
    i = 0
    table_lines = []
    in_table = False
    table_aligns = []

    def flush_table():
        nonlocal table_lines, in_table, table_aligns
        if not table_lines:
            return
        rows = []
        for tl in table_lines:
            if tl.strip().startswith('|') and tl.strip().endswith('|'):
                cells = [c.strip() for c in tl.strip()[1:-1].split('|')]
                rows.append(cells)
        # Filter separator rows
        data_rows = []
        for r in rows:
            if all(re.match(r'^[-:]+$', c) for c in r):
                continue
            data_rows.append(r)
        if not data_rows:
            table_lines = []
            in_table = False
            return

        ncols = max(len(r) for r in data_rows)
        table = doc.add_table(rows=len(data_rows), cols=ncols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_data in enumerate(data_rows):
            for ci, cell_text in enumerate(row_data):
                if ci >= ncols:
                    break
                cell = table.rows[ri].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                if ri == 0:
                    run.bold = True
                    set_cell_shading(cell, 'D9E2F3')
                run.font.size = Pt(9)
                run.font.name = '等线'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

        doc.add_paragraph('')
        table_lines = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        # Table detection
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Headers
        if line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            for run in h.runs:
                run.font.name = '等线'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.name = '等线'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        elif line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            for run in h.runs:
                run.font.name = '等线'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = '等线'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            text = line.strip()[2:]
            # Handle bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = '等线'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        elif line.strip() == '':
            # skip empty in table context, else add small spacer
            pass
        else:
            if line.strip():
                p = doc.add_paragraph()
                # Handle inline bold
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.size = Pt(10)
                    run.font.name = '等线'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        i += 1

    if in_table:
        flush_table()

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '等线'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    title = doc.add_heading('八升九暑假执行计划', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '等线'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('v5 时间块重组版 · 含每日打卡表')
    run.font.size = Pt(14)
    run.font.name = '等线'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026.7.13 - 2026.8.31')
    run.font.size = Pt(12)
    run.font.name = '等线'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    doc.add_page_break()

    # Part 1: 执行计划
    plan_path = BASE / '8周执行计划-v5.md'
    plan_text = plan_path.read_text(encoding='utf-8')
    doc.add_heading('第一部分：八周执行计划', level=1)
    add_markdown_to_doc(doc, plan_text)

    doc.add_page_break()

    # Part 2: 每日打卡表
    checkin_path = BASE / '每日打卡表-v5.md'
    checkin_text = checkin_path.read_text(encoding='utf-8')
    doc.add_heading('第二部分：每日打卡表', level=1)
    add_markdown_to_doc(doc, checkin_text)

    output_path = BASE / '八升九暑假执行计划_v5_含打卡表.docx'
    doc.save(str(output_path))
    print(f'OK -> {output_path}')

if __name__ == '__main__':
    main()
